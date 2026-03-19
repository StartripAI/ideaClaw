"""Document ingester — ingest PDF/DOCX/MD/LaTeX into chunked library.

Supports:
  - PDF (via pdfplumber or PyMuPDF, with fallback to pdfminer)
  - DOCX (via python-docx)
  - Markdown (.md)
  - LaTeX (.tex)

Each document is split into semantic chunks (~500 tokens each)
and stored with metadata for RAG retrieval.

Usage:
    ingester = DocumentIngester(library_dir)
    doc = ingester.ingest("/path/to/paper.pdf")
    print(f"{doc.title}: {len(doc.chunks)} chunks")
"""

from __future__ import annotations

import hashlib
import json
import re
import datetime as dt
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional


@dataclass
class Chunk:
    """A single chunk of document content."""
    id: str
    doc_id: str
    index: int                        # Position in document
    text: str
    section: str = ""                 # Which section this belongs to
    page: int = 0                     # Page number (for PDFs)
    token_count: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, d: Dict[str, Any]) -> "Chunk":
        return cls(**{k: v for k, v in d.items() if k in cls.__dataclass_fields__})


@dataclass
class IngestedDocument:
    """A fully ingested and chunked document."""
    id: str
    title: str
    source_path: str
    file_type: str                    # "pdf", "docx", "md", "tex"
    chunks: List[Chunk] = field(default_factory=list)
    full_text: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    ingested_at: str = ""
    word_count: int = 0

    def to_dict(self) -> Dict[str, Any]:
        d = asdict(self)
        d["chunk_count"] = len(self.chunks)
        # Don't serialize full_text in index (too large)
        d.pop("full_text", None)
        return d


class DocumentIngester:
    """Ingest documents into the personal library.

    Stores at: <library_dir>/documents/
    """

    CHUNK_SIZE = 500       # Target tokens per chunk
    CHUNK_OVERLAP = 50     # Token overlap between chunks

    def __init__(self, library_dir: Optional[Path] = None):
        self.library_dir = library_dir or (Path.home() / ".ideaclaw" / "library")
        self.docs_dir = self.library_dir / "documents"
        self.docs_dir.mkdir(parents=True, exist_ok=True)
        self.index_path = self.library_dir / "doc_index.json"

    def ingest(self, file_path: str | Path) -> IngestedDocument:
        """Ingest a single document.

        Args:
            file_path: Path to PDF, DOCX, MD, or TEX file.

        Returns:
            IngestedDocument with chunks.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        file_type = path.suffix.lower().lstrip(".")
        doc_id = self._make_id(str(path), path.stat().st_mtime)

        # Extract full text
        if file_type == "pdf":
            text, metadata = self._extract_pdf(path)
        elif file_type == "docx":
            text, metadata = self._extract_docx(path)
        elif file_type in ("md", "markdown"):
            text, metadata = self._extract_markdown(path)
        elif file_type in ("tex", "latex"):
            text, metadata = self._extract_latex(path)
        else:
            # Fallback: treat as plain text
            text = path.read_text(encoding="utf-8", errors="ignore")
            metadata = {}
            file_type = "txt"

        # Chunk the text
        chunks = self._chunk_text(text, doc_id)

        doc = IngestedDocument(
            id=doc_id,
            title=metadata.get("title", path.stem),
            source_path=str(path.absolute()),
            file_type=file_type,
            chunks=chunks,
            full_text=text,
            metadata=metadata,
            ingested_at=dt.datetime.now(dt.timezone.utc).isoformat(),
            word_count=len(text.split()),
        )

        # Persist
        self._save_document(doc)
        self._update_index(doc)

        return doc

    def list_documents(self) -> List[Dict[str, Any]]:
        """List all ingested documents."""
        index = self._load_index()
        return index.get("documents", [])

    def get_document(self, doc_id: str) -> Optional[IngestedDocument]:
        """Load a document and its chunks."""
        doc_path = self.docs_dir / f"{doc_id}.json"
        if not doc_path.exists():
            return None
        data = json.loads(doc_path.read_text(encoding="utf-8"))
        chunks = [Chunk.from_dict(c) for c in data.pop("chunks", [])]
        doc = IngestedDocument(**{k: v for k, v in data.items() if k != "chunk_count"})
        doc.chunks = chunks
        return doc

    def get_all_chunks(self) -> List[Chunk]:
        """Load all chunks from all documents (for retrieval index)."""
        all_chunks = []
        for doc_file in self.docs_dir.glob("*.json"):
            try:
                data = json.loads(doc_file.read_text(encoding="utf-8"))
                for c in data.get("chunks", []):
                    all_chunks.append(Chunk.from_dict(c))
            except (json.JSONDecodeError, OSError):
                continue
        return all_chunks

    def remove_document(self, doc_id: str) -> bool:
        """Remove a document from the library."""
        doc_path = self.docs_dir / f"{doc_id}.json"
        if doc_path.exists():
            doc_path.unlink()
            self._remove_from_index(doc_id)
            return True
        return False

    # ---- Text Extraction ----

    def _extract_pdf(self, path: Path) -> tuple[str, Dict]:
        """Extract text from PDF."""
        metadata = {"format": "pdf"}
        text_parts = []

        # Try pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                metadata["pages"] = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
            return "\n\n".join(text_parts), metadata
        except ImportError:
            pass

        # Try PyMuPDF (fitz)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            metadata["pages"] = len(doc)
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n\n".join(text_parts), metadata
        except ImportError:
            pass

        # Fallback: read as binary and try basic extraction
        raw = path.read_bytes()
        text = raw.decode("utf-8", errors="ignore")
        # Strip binary garbage
        text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
        return text, metadata

    def _extract_docx(self, path: Path) -> tuple[str, Dict]:
        """Extract text from DOCX."""
        metadata = {"format": "docx"}
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            metadata["paragraphs"] = len(paragraphs)
            return "\n\n".join(paragraphs), metadata
        except ImportError:
            # Fallback: read XML from docx zip
            import zipfile
            with zipfile.ZipFile(path) as z:
                xml = z.read("word/document.xml").decode("utf-8")
                # Strip XML tags
                text = re.sub(r'<[^>]+>', ' ', xml)
                text = re.sub(r'\s+', ' ', text).strip()
                return text, metadata

    def _extract_markdown(self, path: Path) -> tuple[str, Dict]:
        """Extract text from Markdown."""
        text = path.read_text(encoding="utf-8", errors="ignore")
        sections = re.findall(r'^#{1,3}\s+(.+)', text, re.MULTILINE)
        metadata = {"format": "markdown", "sections": sections}
        return text, metadata

    def _extract_latex(self, path: Path) -> tuple[str, Dict]:
        """Extract text from LaTeX, stripping commands."""
        raw = path.read_text(encoding="utf-8", errors="ignore")
        metadata = {"format": "latex"}

        # Extract title
        title_match = re.search(r'\\title\{([^}]+)\}', raw)
        if title_match:
            metadata["title"] = title_match.group(1)

        # Strip common LaTeX commands but keep content
        text = raw
        # Remove comments
        text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
        # Remove \command{} but keep content inside braces
        text = re.sub(r'\\(?:textbf|textit|emph|text|cite|ref|label)\{([^}]*)\}', r'\1', text)
        # Remove \begin{env}...\end{env} for non-content envs
        text = re.sub(r'\\begin\{(?:figure|table)\}.*?\\end\{(?:figure|table)\}', '', text, flags=re.DOTALL)
        # Remove remaining commands
        text = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})?', '', text)
        # Clean up
        text = re.sub(r'\{|\}', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip(), metadata

    # ---- Chunking ----

    def _chunk_text(self, text: str, doc_id: str) -> List[Chunk]:
        """Split text into overlapping chunks of ~CHUNK_SIZE tokens."""
        words = text.split()
        if not words:
            return []

        chunks = []
        start = 0
        chunk_idx = 0

        while start < len(words):
            end = min(start + self.CHUNK_SIZE, len(words))
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            # Try to detect section from surrounding text
            section = self._detect_section(text, start, words)

            chunk = Chunk(
                id=f"{doc_id}:{chunk_idx}",
                doc_id=doc_id,
                index=chunk_idx,
                text=chunk_text,
                section=section,
                token_count=len(chunk_words),
            )
            chunks.append(chunk)

            # Advance with overlap
            start = end - self.CHUNK_OVERLAP if end < len(words) else end
            chunk_idx += 1

        return chunks

    @staticmethod
    def _detect_section(full_text: str, word_pos: int, words: List[str]) -> str:
        """Try to detect which section a chunk belongs to."""
        # Find approximate character position
        char_pos = len(" ".join(words[:word_pos]))
        text_before = full_text[:char_pos]

        # Find last heading
        headings = list(re.finditer(r'^#{1,3}\s+(.+)|\\section\{([^}]+)\}', text_before, re.MULTILINE))
        if headings:
            match = headings[-1]
            return (match.group(1) or match.group(2) or "").strip()
        return ""

    # ---- Storage ----

    @staticmethod
    def _make_id(path: str, mtime: float) -> str:
        raw = f"{path}:{mtime}"
        return hashlib.sha256(raw.encode()).hexdigest()[:16]

    def _save_document(self, doc: IngestedDocument):
        """Save document + chunks to disk."""
        data = {
            "id": doc.id,
            "title": doc.title,
            "source_path": doc.source_path,
            "file_type": doc.file_type,
            "metadata": doc.metadata,
            "ingested_at": doc.ingested_at,
            "word_count": doc.word_count,
            "chunks": [c.to_dict() for c in doc.chunks],
        }
        doc_path = self.docs_dir / f"{doc.id}.json"
        doc_path.write_text(
            json.dumps(data, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    def _update_index(self, doc: IngestedDocument):
        index = self._load_index()
        docs = index.get("documents", [])
        docs = [d for d in docs if d.get("id") != doc.id]
        docs.insert(0, {
            "id": doc.id,
            "title": doc.title,
            "source_path": doc.source_path,
            "file_type": doc.file_type,
            "chunk_count": len(doc.chunks),
            "word_count": doc.word_count,
            "ingested_at": doc.ingested_at,
        })
        index["documents"] = docs
        index["total_documents"] = len(docs)
        index["updated_at"] = dt.datetime.now(dt.timezone.utc).isoformat()
        self.index_path.write_text(json.dumps(index, indent=2, ensure_ascii=False), encoding="utf-8")

    def _remove_from_index(self, doc_id: str):
        index = self._load_index()
        docs = [d for d in index.get("documents", []) if d.get("id") != doc_id]
        index["documents"] = docs
        index["total_documents"] = len(docs)
        self.index_path.write_text(json.dumps(index, indent=2, ensure_ascii=False), encoding="utf-8")

    def _load_index(self) -> Dict:
        if self.index_path.exists():
            try:
                return json.loads(self.index_path.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                pass
        return {"documents": [], "total_documents": 0}
